#!/usr/bin/env python3
"""Generate Pal Safar product workflow Word and PDF documents."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from fpdf import FPDF

OUT_DIR = Path(__file__).resolve().parent.parent / "docs"
DOCX_PATH = OUT_DIR / "PalSafar-Workflow-Guide.docx"
PDF_PATH = OUT_DIR / "PalSafar-Workflow-Guide.pdf"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xB9, 0x83, 0x4B)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def build_docx():
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Pal Safar")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0xB9, 0x83, 0x4B)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Product Workflow & Features Guide")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.add_run("User • Creator • Vendor Workspaces").font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    add_para(doc, "Explore India, One Spot at a Time — A gamified tourism discovery platform with three distinct workspaces for tourists, content creators, and local businesses.", bold=False)
    doc.add_page_break()

    # TOC-like section list
    add_heading(doc, "Contents", 1)
    toc = [
        "1. App Overview",
        "2. Role Model & Workspace Rules",
        "3. User Workspace — Features & Workflow",
        "4. Creator Workspace — Features & Workflow",
        "5. Vendor Workspace — Features & Workflow",
        "6. Shared Features (All Roles)",
        "7. Pal Points & Rewards System",
        "8. Step-by-Step Key Flows",
        "9. Guest vs Registered User",
        "10. Platform Summary",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()

    # 1. Overview
    add_heading(doc, "1. App Overview", 1)
    add_para(doc, "Pal Safar is a React Native mobile application backed by an Express API and PostgreSQL database. Users explore tourist destinations across India, earn Pal Points through activities, watch and share short video reels, plan trips, and redeem offers at local vendor businesses.")
    add_para(doc, "The app is organized into three primary workspaces:", bold=True)
    add_bullets(doc, [
        "User (Tourist) — Full discovery experience: map, reels, trips, wallet, rewards.",
        "Creator — Content studio for approved creators to publish reels and view analytics.",
        "Vendor — Business portal for tourism businesses to manage offers, scan redemptions, and track analytics.",
    ])
    add_table(doc, ["Component", "Technology"], [
        ["Mobile App", "React Native 0.81"],
        ["Backend API", "Express + TypeScript"],
        ["Database", "PostgreSQL + Prisma"],
        ["Maps", "MapLibre GL Native"],
        ["Media", "Cloudinary"],
        ["Admin", "Next.js Dashboard"],
    ])

    # 2. Role model
    add_heading(doc, "2. Role Model & Workspace Rules", 1)
    add_para(doc, "Every account follows these product rules:")
    add_numbered(doc, [
        "Everyone starts as a User with full tourist app access.",
        "A user may apply once as Creator OR Vendor — not both at the same time.",
        "After admin approval, the profile switcher shows User ↔ one specialty only.",
        "Switching between Creator and Vendor retires the previous specialty role.",
        "Workspace switch updates the server (PATCH /auth/active-mode) and remounts the app navigation.",
    ])
    add_table(doc, ["Mode", "Shell", "Main Tabs"], [
        ["USER", "MainTabs", "Home • Reels • Map • Trips • Profile"],
        ["CONTENT_CREATOR", "CreatorTabs", "Dashboard • Reels • Profile"],
        ["VENDOR", "VendorTabs", "Home • Points • Offers • Analytics • Profile"],
    ])

    # 3. User workspace
    add_heading(doc, "3. User Workspace — Features & Workflow", 1)
    add_heading(doc, "3.1 Features", 2)
    add_table(doc, ["Feature", "Description"], [
        ["Home Dashboard", "Nearby places, quick links to wallet, quests, trips, hidden gems"],
        ["Interactive Map", "MapLibre map with category filters, search, spot markers, vendor pins"],
        ["Geofencing", "50m proximity detection; check-in enabled within 100m of spot"],
        ["Spot Detail", "Place info, photos, reviews, check-in, add to trip"],
        ["Reels Feed", "Vertical short-video feed with like, comment, share"],
        ["Trip Planning", "Manual trip builder and AI itinerary generator"],
        ["Travel Passport", "Visit history, badges, completion stats"],
        ["Wallet", "Pal Points balance and transaction history"],
        ["Rewards", "Browse and redeem vendor offers via QR code"],
        ["Hidden Gems", "Submit community-discovered places for admin review"],
        ["Quests", "Treasure hunts and scavenger challenges"],
        ["Leaderboard", "Regional points rankings"],
        ["Profile", "Journey stats, workspace switcher, settings"],
    ])

    add_heading(doc, "3.2 User Daily Workflow", 2)
    add_numbered(doc, [
        "Open app → Splash → (first time) Onboarding → Login or Sign Up",
        "Land on Home tab — see nearby tourist spots and shortcuts",
        "Open Map tab — explore places; tap marker for Spot Detail",
        "When within range, tap Check In & Claim Points — earn Pal Points",
        "Watch reels on Reels tab; optionally upload if approved as Creator",
        "Plan a trip from Trips tab — manual or AI-generated itinerary",
        "Browse Rewards — redeem Pal Points for vendor discounts (QR shown)",
        "Visit vendor shop — show QR; vendor scans to complete redemption",
        "Submit Hidden Gem from Profile if you discover a new place",
        "Track progress in Travel Passport and Wallet",
    ])

    add_heading(doc, "3.3 User Tab Navigation", 2)
    add_para(doc, "MainTabs (User shell):")
    add_bullets(doc, [
        "Home — Dashboard and discovery entry point",
        "Reels — Community video feed (Explore)",
        "Map — GPS-based exploration",
        "Trips — Itinerary hub and trip management",
        "Profile — Account, journey, settings, workspace switcher",
    ])

    # 4. Creator workspace
    add_heading(doc, "4. Creator Workspace — Features & Workflow", 1)
    add_heading(doc, "4.1 How to Become a Creator", 2)
    add_numbered(doc, [
        "Start as User — use the full tourist app",
        "Apply for Creator status (admin approval required)",
        "Once approved, Profile switcher shows User ↔ Creator",
        "Switch to Creator workspace to access CreatorTabs",
    ])

    add_heading(doc, "4.2 Creator Features", 2)
    add_table(doc, ["Feature", "Description"], [
        ["Creator Dashboard", "Studio overview, stats, quick actions"],
        ["Reels Management", "View and manage your published reels"],
        ["Create Reel", "Record/upload video, add caption, link to tourist spot"],
        ["Daily Bonus", "+100 Pal Points for first reel uploaded each day"],
        ["Creator Profile", "Public creator page with bio and content"],
        ["Analytics", "Views, likes, and engagement metrics"],
        ["Spot Tagging", "Link reels to places for discovery context"],
    ])

    add_heading(doc, "4.3 Creator Reel Upload Workflow", 2)
    add_numbered(doc, [
        "Switch to Creator workspace from Profile",
        "Open Create Reel (from Dashboard or Reels tab)",
        "Select or record video — app compresses video automatically",
        "Add caption, tags, and link to a tourist spot",
        "Tap Post — video uploads to Cloudinary",
        "Server creates reel via POST /social/reels",
        "First reel of the day earns +100 Pal Points",
        "Reel appears in community feed and on linked spot",
    ])

    add_heading(doc, "4.4 Creator Tab Navigation", 2)
    add_bullets(doc, [
        "Dashboard — Studio home and analytics summary",
        "Reels — Your content library and upload entry",
        "Profile — Creator studio settings and public profile",
    ])

    # 5. Vendor workspace
    add_heading(doc, "5. Vendor Workspace — Features & Workflow", 1)
    add_heading(doc, "5.1 How to Become a Vendor", 2)
    add_numbered(doc, [
        "From User Profile, tap Become a Vendor",
        "Complete VendorRegisterScreen — business name, category, location, documents",
        "Submit application — status becomes PENDING",
        "Admin reviews and approves in admin dashboard",
        "Profile switcher shows User ↔ Vendor",
        "Switch to Vendor workspace to manage your business",
    ])

    add_heading(doc, "5.2 Vendor Features", 2)
    add_table(doc, ["Feature", "Description"], [
        ["Vendor Dashboard", "Business overview, offer summary, quick actions"],
        ["Offer Management", "Create, edit, pause, resume discount offers"],
        ["Points Scanner", "Scan tourist QR codes to verify redemptions"],
        ["Redemption History", "Track all verified and pending redemptions"],
        ["Vendor Reels", "Upload promotional videos on business profile"],
        ["Analytics", "Offer views, clicks, redemption metrics"],
        ["Vendor Profile", "Public business page with offers and reviews"],
        ["Pay by Code", "Receive direct Pal Points transfers via vendor code"],
        ["Customer List", "View customers who redeemed at your business"],
    ])

    add_heading(doc, "5.3 Vendor Offer Workflow", 2)
    add_heading(doc, "Create an Offer", 3)
    add_numbered(doc, [
        "Switch to Vendor workspace",
        "Go to Offers tab → Create Offer",
        "Set title, discount type (flat/percentage), points required, limits",
        "Submit — offer goes live after admin approval (if required)",
    ])

    add_heading(doc, "Redemption Flow (Vendor Side)", 3)
    add_numbered(doc, [
        "Tourist redeems offer in app — Pal Points deducted, QR code generated",
        "Tourist shows QR or PAL-XXXXXX token at your shop",
        "Open Points tab (Scanner) or Vendor Redemption screen",
        "Scan QR or enter token manually",
        "App calls POST /redemptions/verify",
        "Confirmation shown — redemption marked VERIFIED",
        "Discount applied per offer terms",
    ])

    add_heading(doc, "5.4 Vendor Tab Navigation", 2)
    add_bullets(doc, [
        "Home — Dashboard and business overview",
        "Points — Scanner and redemption history",
        "Offers — Create and manage discount offers",
        "Analytics — Business performance metrics",
        "Profile — Vendor settings and public page",
    ])

    # 6. Shared features
    add_heading(doc, "6. Shared Features (All Roles)", 1)
    add_table(doc, ["Feature", "Available To"], [
        ["Spot Detail & Search", "All authenticated users"],
        ["Settings & Legal", "All authenticated users"],
        ["Notifications", "All authenticated users"],
        ["Premium Upgrade", "All authenticated users"],
        ["Workspace Switcher", "Users with approved Creator or Vendor role"],
        ["Create Reel screen", "Creator (social API) or Vendor (vendor reel API)"],
        ["Offline banner", "All — syncs when back online"],
    ])

    # 7. Pal Points
    add_heading(doc, "7. Pal Points & Rewards System", 1)
    add_para(doc, "Pal Points are the in-app currency earned through tourism activities and spent at vendor businesses.")
    add_heading(doc, "Earn Points", 2)
    add_table(doc, ["Action", "How"], [
        ["Check in at spot", "Visit place within range, tap Check In"],
        ["Write review", "Submit review on Spot Detail"],
        ["Upload reel", "Creator: +100 pts first reel/day"],
        ["Hidden gem approved", "Community submission approved by admin"],
        ["Complete quest", "Finish treasure hunt or challenge"],
        ["Campaign claim", "Participate in reward campaigns"],
    ])
    add_heading(doc, "Spend Points", 2)
    add_bullets(doc, [
        "Redeem vendor offers — QR generated, points deducted at generation",
        "Pay vendor directly — enter vendor code and point amount",
    ])

    # 8. Key flows
    add_heading(doc, "8. Step-by-Step Key Flows", 1)

    add_heading(doc, "8.1 Sign Up Flow", 2)
    add_numbered(doc, [
        "Open app → Onboarding (first launch) → Sign Up",
        "Enter name, email, password",
        "Server creates account + wallet (0 points)",
        "JWT stored securely → User lands on MainTabs",
    ])

    add_heading(doc, "8.2 Check-In Flow", 2)
    add_numbered(doc, [
        "Open Map → navigate near tourist spot",
        "Tap marker → Spot Detail opens",
        "App calculates GPS distance",
        "Within 100m → Check In button enabled",
        "Tap Check In → points awarded server-side",
        "Passport and wallet updated",
    ])

    add_heading(doc, "8.3 Offer Redemption Flow (End-to-End)", 2)
    add_numbered(doc, [
        "User browses Rewards / Vendor Offers",
        "Selects offer → Redeem",
        "Points deducted → QR code + PAL-token shown (15 min expiry)",
        "User visits vendor location",
        "Vendor scans token → verified",
        "Transaction complete — both see confirmation",
    ])

    add_heading(doc, "8.4 Workspace Switch Flow", 2)
    add_numbered(doc, [
        "Open Profile → Workspace Switcher",
        "Tap User, Creator, or Vendor",
        "Server updates activeMode",
        "App remounts navigation to correct tab shell",
        "User continues in new workspace",
    ])

    add_heading(doc, "8.5 Hidden Gem Flow", 2)
    add_numbered(doc, [
        "User discovers unlisted place",
        "Profile → Add Hidden Gem",
        "Fill details, GPS, photos → Submit",
        "Status: Pending",
        "Admin approves → place appears on map for everyone",
        "Submitter may earn bonus points",
    ])

    # 9. Guest
    add_heading(doc, "9. Guest vs Registered User", 1)
    add_table(doc, ["Action", "Guest", "Registered User"], [
        ["Browse map & places", "Yes", "Yes"],
        ["Watch reels", "Yes", "Yes"],
        ["Check in", "No", "Yes"],
        ["Redeem offers", "No", "Yes"],
        ["Write reviews", "No", "Yes"],
        ["Submit hidden gems", "No", "Yes"],
        ["Upload reels", "No", "Yes (if Creator/Vendor)"],
        ["Wallet & points", "No", "Yes"],
        ["Workspace switch", "No", "Yes (if approved)"],
    ])

    # 10. Summary
    add_heading(doc, "10. Platform Summary", 1)
    add_para(doc, "Pal Safar connects three audiences in one platform:")
    add_bullets(doc, [
        "Users explore India, earn points, plan trips, and redeem local offers.",
        "Creators publish spot-linked reels and build an audience.",
        "Vendors attract tourists with point-based offers and verify redemptions in-app.",
    ])
    add_para(doc, "All workspaces share one account and switch seamlessly via the Profile switcher after role approval.")
    add_para(doc, "")
    add_para(doc, "Document generated for Pal Safar • Explore India, One Spot at a Time", bold=True)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"Created: {DOCX_PATH}")


def build_pdf():
    """Generate PDF with fpdf2 (Latin-safe encoding)."""
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    w = pdf.epw  # effective page width

    def section(title, size=14):
        pdf.set_font("Helvetica", "B", size)
        pdf.set_text_color(185, 131, 75)
        pdf.multi_cell(w, 8, title)
        pdf.ln(2)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 11)

    def body(text):
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(w, 6, text)
        pdf.ln(2)

    def bullets(items):
        pdf.set_font("Helvetica", "", 11)
        for item in items:
            pdf.multi_cell(w, 6, f"- {item}")
        pdf.ln(2)

    def numbered(items):
        pdf.set_font("Helvetica", "", 11)
        for i, item in enumerate(items, 1):
            pdf.multi_cell(w, 6, f"{i}. {item}")
        pdf.ln(2)

    # Title
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(185, 131, 75)
    pdf.cell(w, 12, "Pal Safar", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 14)
    pdf.set_text_color(139, 115, 85)
    pdf.cell(w, 8, "Product Workflow & Features Guide", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(w, 8, "User  |  Creator  |  Vendor", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)
    body("Explore India, One Spot at a Time - A gamified tourism discovery platform with three workspaces for tourists, content creators, and local businesses.")

    section("1. App Overview")
    body("Pal Safar is a mobile tourism app where users explore destinations, earn Pal Points, share reels, plan trips, and redeem vendor offers.")
    bullets([
        "User (Tourist): Map, reels, trips, wallet, rewards, hidden gems",
        "Creator: Reel studio, analytics, daily point bonus",
        "Vendor: Offers, QR scanner, business analytics",
    ])

    section("2. Role Model")
    numbered([
        "Everyone starts as User with full tourist access",
        "Apply once as Creator OR Vendor (not both)",
        "After approval: switch User and one specialty only",
        "Switching retires the other specialty role",
    ])

    section("3. User Workspace")
    body("Tabs: Home | Reels | Map | Trips | Profile")
    bullets([
        "Home: Nearby places, wallet and quest shortcuts",
        "Map: GPS exploration, check-in within 100m",
        "Reels: Watch community videos",
        "Trips: Manual and AI trip planning",
        "Profile: Passport, settings, workspace switcher",
    ])
    section("User Workflow", 12)
    numbered([
        "Sign up and land on Home",
        "Explore Map and check in at spots to earn points",
        "Watch reels and plan trips",
        "Redeem offers from Rewards - show QR at vendor",
        "Submit hidden gems for community discovery",
    ])

    section("4. Creator Workspace")
    body("Tabs: Dashboard | Reels | Profile")
    bullets([
        "Requires admin-approved Creator role",
        "Upload reels linked to tourist spots",
        "+100 Pal Points for first reel each day",
        "View analytics and manage creator profile",
    ])
    section("Creator Reel Workflow", 12)
    numbered([
        "Switch to Creator workspace from Profile",
        "Open Create Reel, select video and spot",
        "Video compresses and uploads to Cloudinary",
        "Reel published to community feed",
    ])

    section("5. Vendor Workspace")
    body("Tabs: Home | Points | Offers | Analytics | Profile")
    bullets([
        "Register business - admin approval required",
        "Create discount offers (points required)",
        "Scan tourist QR codes to verify redemptions",
        "Upload vendor reels and view analytics",
        "Receive direct point payments via vendor code",
    ])
    section("Vendor Redemption Workflow", 12)
    numbered([
        "Create and publish offers",
        "Tourist redeems in app - QR generated",
        "Scan QR or enter PAL-token at Points tab",
        "Redemption verified - discount applied",
    ])

    section("6. Pal Points System")
    body("Earn: check-in, reviews, reels, quests, approved hidden gems")
    body("Spend: redeem offers (QR) or pay vendor by code")

    section("7. Guest vs Registered")
    body("Guests can browse map and reels only. Registered users get check-in, wallet, redemption, reviews, and role workspaces.")

    section("8. Platform Summary")
    body("Users explore and earn. Creators publish content. Vendors convert points into foot traffic. One account, three workspaces.")

    pdf.ln(5)
    pdf.set_font("Helvetica", "I", 10)
    pdf.cell(w, 6, "Pal Safar - Explore India, One Spot at a Time", new_x="LMARGIN", new_y="NEXT", align="C")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf.output(str(PDF_PATH))
    print(f"Created: {PDF_PATH}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
